"""Document processing: text extraction, OCR for scanned pages, chunking.

Supported: PDF (native text + OCR fallback for scanned pages), images,
CSV/XLSX, DOCX. All local via PyMuPDF / Pillow / PaddleOCR.
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path

from PIL import Image

logger = logging.getLogger("nexus.docproc")

ALLOWED_EXTS = {".pdf", ".png", ".jpg", ".jpeg", ".csv", ".xlsx", ".docx"}


def extract_text(path: Path, ext: str) -> tuple[str, int, dict]:
    """Return (text, page_count, extra_metadata)."""
    ext = ext.lower().lstrip(".")
    if ext in ("pdf",):
        return _extract_pdf(path)
    if ext in ("png", "jpg", "jpeg"):
        img = Image.open(path)
        from app.services.ocr_service import ocr_image
        text = ocr_image(img.convert("RGB"))
        return text, 1, {"engine": "ocr"}
    if ext == "csv":
        txt = path.read_text(encoding="utf-8", errors="replace")
        return txt, 1, {"rows": txt.count("\n")}
    if ext == "xlsx":
        return _extract_xlsx(path)
    if ext == "docx":
        return _extract_docx(path)
    return "", 0, {}


def _extract_xlsx(path: Path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts, rows = [], 0
    for ws in wb.worksheets:
        parts.append(f"# SHEET: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            rows += 1
            vals = ["" if c is None else str(c) for c in row]
            if any(v.strip() for v in vals):
                parts.append(" | ".join(vals))
    return "\n".join(parts), wb.sheetnames.__len__(), {"rows": rows}


def _extract_docx(path: Path):
    from docx import Document
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts), 1, {"paragraphs": len(parts)}


def _extract_pdf(path: Path) -> tuple[str, int, dict]:
    import fitz  # pymupdf
    doc = fitz.open(str(path))
    pages = doc.page_count
    parts: list[str] = []
    meta = {"scanned": False, "text_pages": 0, "ocr_pages": 0}
    from app.services.ocr_service import available_engine, ocr_image

    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text and text.strip():
            parts.append(f"--- PAGE {i + 1} ---\n{text}")
            meta["text_pages"] += 1
            continue
        # scanned page -> render + OCR
        if available_engine() != "unavailable":
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            try:
                ocr = ocr_image(img)
                if ocr.strip():
                    parts.append(f"--- PAGE {i + 1} (OCR) ---\n{ocr}")
                    meta["ocr_pages"] += 1
                    meta["scanned"] = True
            except Exception as e:  # noqa
                logger.warning("page %d OCR failed: %s", i + 1, e)
        else:
            meta["scanned"] = True
    return "\n\n".join(parts), pages, meta


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[str]:
    """Split text into overlapping, paragraph-aware chunks."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buffer = ""
    for p in paragraphs:
        if len(buffer) + len(p) + 2 <= chunk_size:
            buffer = (buffer + "\n\n" + p).strip()
        else:
            if buffer:
                chunks.append(buffer)
            # split oversized paragraph by sentences
            if len(p) > chunk_size:
                sentences = re.split(r"(?<=[.!?])\s+", p)
                cur = ""
                for s in sentences:
                    if len(cur) + len(s) + 1 <= chunk_size:
                        cur = (cur + " " + s).strip()
                    else:
                        if cur:
                            chunks.append(cur)
                        cur = s
                if cur:
                    buffer = cur
            else:
                buffer = p
    if buffer:
        chunks.append(buffer)
    return chunks
